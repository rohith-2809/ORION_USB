import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

class DocumentWriter:
    """
    ORION Docx Engine — Final Optimized Version
    ------------------------------------------
    • Support for Markdown headers (#, ##, ###)
    • Support for bold (**text**) and italics (*text*)
    • Support for lists (-, *, 1.)
    • Academic-ready typography
    """

    @staticmethod
    def write_docx(path: str, content: str):
        doc = Document()
        DocumentWriter._configure_styles(doc)

        lines = content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # ───── HEADERS ─────
            if line.startswith("#"):
                level = len(line.split()[0])
                text = line.lstrip("#").strip()
                # Strip internal markdown (** or __) from headers for clean look
                text = text.replace("**", "").replace("__", "")
                style = "ORION_Title" if level == 1 else f"ORION_Heading{min(level, 3)}"
                doc.add_paragraph(text, style=style)
                continue

            # ───── HORIZONTAL RULE ─────
            if line.startswith("---") or line.startswith("===") or line.startswith("___"):
                p = doc.add_paragraph()
                run = p.add_run("__________________________________________________")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(200, 200, 200)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # ───── LISTS ─────
            if re.match(r"^[-*]\s+", line):
                text = re.sub(r"^[-*]\s+", "", line)
                p = doc.add_paragraph(style="ORION_Bullet")
                DocumentWriter._add_rich_text(p, text)
                continue

            if re.match(r"^\d+\.\s+", line):
                text = re.sub(r"^\d+\.\s+", "", line)
                p = doc.add_paragraph(style="ORION_ListNumber")
                DocumentWriter._add_rich_text(p, text)
                continue

            # ───── STANDARD PARAGRAPH ─────
            p = doc.add_paragraph(style="ORION_Body")
            DocumentWriter._add_rich_text(p, line)

        doc.save(path)

    # ──────────────────────────────────────────
    # TEXT PARSER (BOLD/ITALIC)
    # ──────────────────────────────────────────
    @staticmethod
    def _add_rich_text(paragraph, text):
        # Regex to split by **bold** or *italic*
        # This splits into: ['normal ', '**bold**', ' normal ', '*italic*', '']
        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)

        for token in tokens:
            if not token: continue

            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                # Handle bullet points that might be mistaken for italics if incomplete?
                # No, regex above ensures matching pairs.
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)

    # ──────────────────────────────────────────
    # STYLE CONFIGURATION
    # ──────────────────────────────────────────
    @staticmethod
    def _configure_styles(doc: Document):
        styles = doc.styles

        def safe_style(name, style_type):
            return styles[name] if name in styles else styles.add_style(name, style_type)

        # ───── TITLE (Header 1) ─────
        title = safe_style("ORION_Title", WD_STYLE_TYPE.PARAGRAPH)
        title.font.name = "Times New Roman"
        title.font.size = Pt(24)
        title.font.bold = True
        title.font.color.rgb = RGBColor(0, 0, 0)
        title.paragraph_format.space_after = Pt(12)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ───── HEADING 1 (Markdown ##) ─────
        h1 = safe_style("ORION_Heading1", WD_STYLE_TYPE.PARAGRAPH)
        h1.font.name = "Times New Roman"
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(46, 116, 181) # Orion Blue
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(6)

        # ───── HEADING 2 (Markdown ###) ─────
        h2 = safe_style("ORION_Heading2", WD_STYLE_TYPE.PARAGRAPH)
        h2.font.name = "Times New Roman"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(68, 84, 106) # Slate
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)

         # ───── HEADING 3 (Markdown ####) ─────
        h3 = safe_style("ORION_Heading3", WD_STYLE_TYPE.PARAGRAPH)
        h3.font.name = "Times New Roman"
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.italic = True
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(0)

        # ───── BODY ─────
        body = safe_style("ORION_Body", WD_STYLE_TYPE.PARAGRAPH)
        body.font.name = "Times New Roman"
        body.font.size = Pt(11)
        body.paragraph_format.space_after = Pt(8)
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ───── BULLET ─────
        bullet = safe_style("ORION_Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.font.name = "Times New Roman"
        bullet.font.size = Pt(11)
        bullet.paragraph_format.left_indent = Inches(0.25)
        bullet.paragraph_format.space_after = Pt(2)

        # ───── NUMBERED LIST ─────
        numlist = safe_style("ORION_ListNumber", WD_STYLE_TYPE.PARAGRAPH)
        numlist.font.name = "Times New Roman"
        numlist.font.size = Pt(11)
        numlist.paragraph_format.left_indent = Inches(0.25)
        numlist.paragraph_format.space_after = Pt(2)
